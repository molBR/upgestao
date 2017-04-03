from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING

document = Document('GU.docx')

nome = input ("Digite o nome completo: ")
email = input ("Digite o email: ")
tipoevento = input ("Tipo evento: ")
tipocardapio = input ("tipo de cardapio: ")
data = input ("data: ")
numeropessoas = input ("numero de pessoas: ")
local = input ("Local: ")
horario = input ("Horario: ")

p = document.add_paragraph()
p.add_run('Nome completo: ').bold = True
p.add_run(nome)
p.add_run('\n')
p.add_run('E-mail: ').bold = True
p.add_run(email)
p.add_run('\n')
p.add_run('Tipo do evento: ').bold = True
p.add_run(tipoevento)
p.add_run('\n')
p.add_run('Tipo do cardápio: ').bold = True
p.add_run(tipocardapio)
p.add_run('\n')
p.add_run('Data do evento: ').bold = True
p.add_run(data)
p.add_run('\n')
p.add_run('Número de pessoas: ').bold = True
p.add_run(numeropessoas)
p.add_run('\n')
p.add_run('Local do evento: ').bold = True
p.add_run(local)
p.add_run('\n')
p.add_run('Horário: ').bold = True
p.add_run(horario).bold = True
p.add_run('\n')

cont = 0

table = document.add_table(rows=4, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Salgado'
hdr_cells[1].text = 'Quantidade'
hdr_cells[2].text = 'Preço'

cont=5
table = document.add_table(rows=5, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Bebidas'

cont=cont+5
table = document.add_table(rows=4, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Pessoal'

cont=cont+5
table = document.add_table(rows=1, cols=1)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Valor Total:'
cont= cont+1




rodape = document.add_paragraph()
rodape.paragraph_format.keep_together = True
rodape.paragraph_format.keep_with_next = True

run = rodape.add_run('Materiais: \n')
run.bold = True
run.underline = True
run.font.size = Pt (11)
rodape.add_run('Bandejas e talheres em aço inox, pratos em porcelana, copos em vidro, toalhas de mesa,'
               ' cadeiras, gelo para conservar bebidas geladas ou freezer, todo material de cozinha e '
               'todos os demais matérias necessários para a execução do serviço solicitado acima.\n')
run = rodape.add_run('Forma de pagamento: ')
run.bold = True
run.underline = True
run.font.size = Pt (11)
rodape.add_run('10% na assinatura do contrato e o restante 15 dias antes do evento.\n').bold = True
rodape.add_run('Obs. 1:Em caso de perda ou dano dos materias utilizados, será cobrado o valor do mesmo:\n'
               '*Copos de cerveja, refrigerante e água(unid.)- 5,00 *Taças para vinho(unid.)-8,50\n'
               'Obs. 2: A duração do evento é de 5 horas, contadas a partir do horário marcado.\n'
               'Excedendo será cobrado uma taxa de R$150,00 por hora, mais R$ 40,00 por mão de obra disponível.\n ')
run = rodape.add_run('Obs. 3:Todas as bebidas terào abatimento no preço de custo.\n'
                     'Todo cancelamento ou quebra de contrato acarretará na multa de 5% do valor total do orçamento.\n');
run.bold = True
run.underline = True
run.font.size = Pt (11)

run = rodape.add_run('Orçamento válido por 30 dias. ')
run.bold = True
run.font.size = Pt (11)





nome1= nome + '.docx'
nome1_stream = nome1
document.save(nome1_stream)
