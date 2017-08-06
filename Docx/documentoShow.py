from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING



def docx(ClienteInfo,VendaInfo,ProdInfo,ItensSalgados,itensBebidas,ItensPessoal):

    print ProdInfo
    print ClienteInfo
    print VendaInfo

    docesL = []
    salgadosL = []
    massasL = []
    bebidasL = []
    outrosL = []

    for i in range(len(ProdInfo)):
        if ProdInfo[i][5] == "1":
            docesL.append(ProdInfo[i])
        if ProdInfo[i][5] == "2":
            salgadosL.append(ProdInfo[i])
        if ProdInfo[i][5] == "3":
            massasL.append(ProdInfo[i])
        if ProdInfo[i][5] == "4":
            bebidasL.append(ProdInfo[i])
        if ProdInfo[i][5] == "5":
            outrosL.append(ProdInfo[i])
    document = Document('Docx\GU.docx')
    #document = Document('GU.docx')
    nome = ClienteInfo[1]
    email = ClienteInfo[5]
    telefone = ClienteInfo[4]
    tipoevento = VendaInfo[3]
    data = VendaInfo[7]
    numeroAdulto = VendaInfo[4]
    numeroCrianca = VendaInfo[5]
    local = VendaInfo[6]
    p = document.add_paragraph()
    p.add_run('Nome completo: ').bold = True
    p.add_run(nome)
    p.add_run('\n')
    p.add_run('E-mail: ').bold = True
    p.add_run(email)
    p.add_run('Telefone: ').bold = True
    p.add_run(telefone)
    p.add_run('\n')
    p.add_run('Tipo do evento: ').bold = True
    p.add_run(tipoevento)
    p.add_run('\n')
    p.add_run('Data do evento: ').bold = True
    p.add_run(data)
    p.add_run('\n')
    p.add_run('Numero de Adultos: ').bold = True
    p.add_run(numeroAdulto)
    p.add_run('\n')
    p.add_run('Numero de Criancas: ').bold = True
    p.add_run(numeroCrianca)
    p.add_run('\n')
    p.add_run('Local do evento: ').bold = True
    p.add_run(local)
    p.add_run('\n')


    salgados = [['kibe', '100', '123', ],
                ['batata frita', '35', '19'],
                ['pao de queijo', '100', '621']]

    cont = 0
    #ItensSalgados = 3
    #itensBebidas = 3
    #ItensPessoal = 3

    Itens = ItensSalgados + itensBebidas + ItensPessoal
    TotalLinha = 4 + Itens + (int((Itens - 19) / 27) + (1 if Itens > 19 else 0))

    def titulo(TituloSalgado):
        TituloSalgado[0].text = 'SALGADO'
        TituloSalgado[1].text = 'QUANTIDADE'
        TituloSalgado[2].text = 'PREcO'

    table = document.add_table(rows=TotalLinha, cols=3)

    contLinha = 0
    contSalgado = 0
    while contSalgado < ItensSalgados:
        TabelaSalgados = table.rows[contLinha].cells
        if contLinha == 0 or contLinha == 20 or contLinha % 27 - 20 == 0:
            titulo(TabelaSalgados)
        else:
            TabelaSalgados[0].text = salgados[0][0]
            TabelaSalgados[1].text = salgados[0][1]
            TabelaSalgados[2].text = salgados[0][2]
            contSalgado = contSalgado + 1

        contLinha = contLinha + 1

    TituloBebidas = table.rows[contLinha].cells
    TituloBebidas[0].text = 'BEBIDAS'
    TituloBebidas[1].text = 'QUANTIDADE'
    TituloBebidas[2].text = 'PREcO'
    contLinha = contLinha + 1
    contBebidas = 0
    bebidas = [['coca-cola', '100', '123', ],
               ['batata', '35', '19']]

    while contBebidas < itensBebidas:
        TabelaBebidas = table.rows[contLinha].cells
        if contLinha == 0 or contLinha == 20 or contLinha % 27 - 20 == 0:
            TituloBebidas = table.rows[contLinha].cells
            TituloBebidas[0].text = 'BEBIDAS'
            TituloBebidas[1].text = 'QUANTIDADE'
            TituloBebidas[2].text = 'PREcO'
            contLinha = contLinha
        else:
            TabelaBebidas[0].text = bebidas[0][0]
            TabelaBebidas[1].text = bebidas[0][1]
            TabelaBebidas[2].text = bebidas[0][2]
            contBebidas = contBebidas + 1

        contLinha = contLinha + 1

    TituloPessoal = table.rows[contLinha].cells
    TituloPessoal[0].text = 'PESSOAL'
    TituloPessoal[1].text = 'QUANTIDADE'
    TituloPessoal[2].text = 'PREcO'
    contLinha = contLinha + 1
    contPessoal = 0
    pessoal = [['seguranca', '100', '123', ],
               ['aa', '35', '19'],
               'sa','2','23']

    while contPessoal < ItensPessoal:
        TabelaPessoal = table.rows[contLinha].cells
        if contLinha == 0 or contLinha == 20 or contLinha % 27 - 20 == 0:
            TituloPessoal = table.rows[contLinha].cells
            TituloPessoal[0].text = 'PESSOAL'
            TituloPessoal[1].text = 'QUANTIDADE'
            TituloPessoal[2].text = 'PREcO'
            contLinha = contLinha
        else:
            TabelaPessoal[0].text = pessoal[0][0]
            TabelaPessoal[1].text = pessoal[0][1]
            TabelaPessoal[2].text = pessoal[0][2]
            contPessoal = contPessoal + 1

        contLinha = contLinha + 1

    table = document.add_table(rows=1, cols=1)
    TabelaTotal = table.rows[0].cells
    TabelaTotal[0].text = 'Valor Total:'
    cont = cont + 1

    print(cont)
    rodape = document.add_paragraph()
    rodape.paragraph_format.keep_together = True
    rodape.paragraph_format.keep_with_next = True

    run = rodape.add_run('Materiais: \n')
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    rodape.add_run('Bandejas e talheres em aco inox, pratos em porcelana, copos em vidro, toalhas de mesa,'
                   ' cadeiras, gelo para conservar bebidas geladas ou freezer, todo material de cozinha e '
                   'todos os demais materias necessarios para a execucao do servico solicitado acima.\n')
    run = rodape.add_run('Forma de pagamento: ')
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    rodape.add_run('10% na assinatura do contrato e o restante 15 dias antes do evento.\n').bold = True
    rodape.add_run('Obs. 1:Em caso de perda ou dano dos materias utilizados, sera cobrado o valor do mesmo:\n'
                   '*Copos de cerveja, refrigerante e agua(unid.)- 5,00 *Tacas para vinho(unid.)-8,50\n'
                   'Obs. 2: A duracao do evento e de 5 horas, contadas a partir do horario marcado.\n'
                   'Excedendo sera cobrado uma taxa de R$150,00 por hora, mais R$ 40,00 por mao de obra disponivel.\n ')
    run = rodape.add_run('Obs. 3:Todas as bebidas terao abatimento no preco de custo.\n'
                         'Todo cancelamento ou quebra de contrato acarretara na multa de 5% do valor total do orcamento.\n');
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)

    run = rodape.add_run('Orcamento valido por 30 dias. ')
    run.bold = True
    run.font.size = Pt(11)

    nome1 = nome + '.docx'
    nome1_stream = nome1
    print nome1_stream
    document.save(nome1_stream)
    return 0
#It = 2
#docx(It,It,It)